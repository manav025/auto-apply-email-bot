"""
Job Application Email Assistant -- Streamlit web app version.

Same functionality as job_email_sender.py, but as a browser UI you can run
locally (`streamlit run streamlit_app.py`) or host for free on Streamlit
Community Cloud, so nothing runs on your own laptop.

All credentials (Groq, Gmail, Adzuna) come from Streamlit secrets only --
there is no UI field to type them in, so they never appear on screen or in
browser session state. Set them in .streamlit/secrets.toml locally, or in
your app's Settings -> Secrets on Streamlit Community Cloud.
"""

import io
import re
import time
import json

import pandas as pd
import pdfplumber
import docx
import requests
import streamlit as st
from groq import Groq
import smtplib
import ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

st.set_page_config(page_title="CV Analyzer & Job Finder", page_icon="📄", layout="wide")

st.markdown("""
<style>
:root {
  --bg: #f7f6fd; --bg2: #efedfb; --card-bg: #ffffff; --text: #1e1b4b; --text-muted: #6b7280;
  --border: #ece9f9; --input-bg: #ffffff; --purple: #6D5BD0; --track: #ece9f9;
}
@media (prefers-color-scheme: dark) {
  :root {
    --bg: #14122a; --bg2: #1a1733; --card-bg: #1f1c3d; --text: #f1eefc; --text-muted: #b3aed6;
    --border: #34305f; --input-bg: #262247; --purple: #a996ff; --track: #322d5c;
  }
}
.stApp { background: linear-gradient(180deg, var(--bg) 0%, var(--bg2) 100%) !important; }
.stApp, .stApp p, .stApp span, .stApp label, .stMarkdown { color: var(--text); }
.score-card, .job-card {
  background: var(--card-bg) !important; color: var(--text);
  border-radius:20px; padding:24px; box-shadow:0 2px 14px rgba(76,63,176,0.10);
  border:1px solid var(--border);
}
.gauge-wrap { display:flex; flex-direction:column; align-items:center; padding-top:8px; }
.chip { display:inline-block; padding:6px 14px; border-radius:999px; font-weight:600; font-size:0.85rem; margin:4px 6px 4px 0; }
.chip-green { background:#0f5132; color:#8ef7c1; }
.chip-red { background:#5c1a1a; color:#ffb4b4; }
@media (prefers-color-scheme: light) {
  .chip-green { background:#d1fae5; color:#047857; }
  .chip-red { background:#fee2e2; color:#b91c1c; }
}
.bar-row { margin-bottom:16px; }
.bar-label { display:flex; justify-content:space-between; font-weight:600; color:var(--text); margin-bottom:5px; }
.bar-track { background:var(--track); border-radius:8px; height:9px; overflow:hidden; }
.bar-fill { height:100%; border-radius:8px; }
.badge { display:inline-block; padding:3px 11px; border-radius:999px; font-size:0.72rem; font-weight:700; background:#3a3166; color:#cabcff; margin-left:6px; vertical-align:middle; }
.badge-ok { background:#0f5132; color:#8ef7c1; }
.badge-warn { background:#5c4a12; color:#ffe08a; }
@media (prefers-color-scheme: light) {
  .badge { background:#ede9fe; color:#5b21b6; }
  .badge-ok { background:#d1fae5; color:#047857; }
  .badge-warn { background:#fef3c7; color:#92400e; }
}
.section-title { font-size:1.4rem; font-weight:800; color:var(--text); margin:0.2em 0 0.4em 0; }
.skill-chip { display:inline-block; padding:5px 12px; border-radius:999px; font-size:0.82rem; font-weight:600;
  margin:3px 5px 3px 0; background:var(--track); color:var(--purple); border:1px solid var(--border); }
/* Nudge common native widgets to follow the same palette so they don't clash */
[data-testid="stFileUploaderDropzone"], .stTextInput input, .stTextArea textarea,
[data-baseweb="select"], .stNumberInput input {
  background: var(--input-bg) !important; color: var(--text) !important; border-color: var(--border) !important;
}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Optional password gate -- set APP_PASSWORD in Streamlit secrets to enable.
# ---------------------------------------------------------------------------
_app_password = st.secrets.get("APP_PASSWORD", "")
if _app_password:
    if not st.session_state.get("unlocked"):
        st.title("🔒 CV Analyzer & Job Finder")
        pw = st.text_input("App password", type="password")
        if st.button("Unlock"):
            if pw == _app_password:
                st.session_state.unlocked = True
                st.rerun()
            else:
                st.error("Wrong password")
        st.stop()

st.title("📄 CV Analyzer & Job Finder")
st.caption("ATS scoring, JD match, AI-tailored resume, and outreach emails.")

# ---------------------------------------------------------------------------
# Credentials -- backend only, via Streamlit secrets. No UI fields for these,
# so they're never typed, shown, or stored in the browser. Configure them in
# .streamlit/secrets.toml (local) or Settings -> Secrets (Streamlit Cloud).
# ---------------------------------------------------------------------------
groq_key = st.secrets.get("GROQ_API_KEY", "")
groq_model = st.secrets.get("GROQ_MODEL", "llama-3.3-70b-versatile")
gmail_address = st.secrets.get("GMAIL_ADDRESS", "")
gmail_app_password = st.secrets.get("GMAIL_APP_PASSWORD", "")
sender_name = st.secrets.get("SENDER_NAME", "")
adzuna_app_id = st.secrets.get("ADZUNA_APP_ID", "")
adzuna_app_key = st.secrets.get("ADZUNA_APP_KEY", "")

with st.sidebar:
    st.header("Setup status")
    st.caption("All credentials live in Streamlit secrets, not on this page.")


    def status_row(label, ok):
        icon = "✅" if ok else "⬜"
        st.markdown(f"{icon} {label}")


    status_row("Groq API key (required)", bool(groq_key))
    status_row("Gmail (for sending)", bool(gmail_address and gmail_app_password))
    status_row("Adzuna (for job search)", bool(adzuna_app_id and adzuna_app_key))
    if not groq_key:
        st.warning("Add `GROQ_API_KEY` to `.streamlit/secrets.toml` to enable AI features. "
                   "Free key: console.groq.com/keys")
    st.markdown("---")
    seconds_between = st.slider("Seconds to wait between sends", 5, 90, 30)
    st.markdown("---")
    st.markdown("[Free Groq key](https://console.groq.com/keys) · "
                "[Gmail app passwords](https://myaccount.google.com/apppasswords) · "
                "[Free Adzuna API key](https://developer.adzuna.com/)")


def call_ai(prompt: str, max_tokens: int = 1000) -> str:
    if not groq_key:
        raise RuntimeError("Enter your Groq API key in the sidebar first.")
    client = Groq(api_key=groq_key)
    resp = client.chat.completions.create(
        model=groq_model,
        max_tokens=max_tokens,
        temperature=0.6,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.choices[0].message.content.strip()


def extract_resume_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        text = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    elif name.endswith(".docx"):
        d = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in d.paragraphs)
    elif name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    else:
        raise ValueError("Unsupported resume format")


def send_email(to_addr: str, subject: str, body: str, resume_bytes: bytes, resume_name: str):
    msg = MIMEMultipart()
    msg["From"] = f"{sender_name} <{gmail_address}>" if sender_name else gmail_address
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    part = MIMEApplication(resume_bytes, Name=resume_name)
    part["Content-Disposition"] = f'attachment; filename="{resume_name}"'
    msg.attach(part)

    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(gmail_address, gmail_app_password)
        server.sendmail(gmail_address, to_addr, msg.as_string())


def name_from_email(email: str) -> str:
    """Best-effort guess at a person's name from their email address, e.g.
    'jane.doe@acme.com' -> 'Jane Doe', 'jdoe123@acme.com' -> 'Jdoe123'.
    Used only as a fallback when no name was supplied."""
    if not email or "@" not in email:
        return ""
    local = email.split("@")[0]
    local = re.sub(r"\d+", "", local)  # drop trailing numbers like jdoe123
    parts = re.split(r"[._\-+]", local)
    parts = [p for p in parts if p]
    if not parts:
        return ""
    return " ".join(p.capitalize() for p in parts)


CONTACT_OPTIONAL_COLUMNS = [
    "name",
    "role",
    "company",
    "job_title",
    "company_note",
    "contact_note",
    "project_match",
    "tone",
    "custom_subject",
]
CONTACT_REQUIRED_COLUMNS = ["email"]
CONTACT_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
GENERIC_DRAFT_PATTERNS = [
    r"\bi am writing to express\b",
    r"\bto whom it may concern\b",
    r"\bdear hiring manager\b",
    r"\bplease find attached\b",
    r"\bany opportunities\b",
]


def validate_uploaded_contacts(df: pd.DataFrame) -> tuple[list[int], list[int], list[str]]:
    email_series = df["email"].astype(str).str.strip()
    blank_rows = (email_series == "")
    invalid_rows = (~blank_rows) & (~email_series.str.match(CONTACT_EMAIL_RE))
    row_numbers_blank = (df.index[blank_rows] + 2).tolist()
    row_numbers_invalid = (df.index[invalid_rows] + 2).tolist()
    dup_emails = email_series[email_series != ""].str.lower()
    duplicate_values = sorted(dup_emails[dup_emails.duplicated(keep=False)].unique().tolist())
    return row_numbers_blank, row_numbers_invalid, duplicate_values


def evaluate_draft_quality(body: str, contact_name: str, company: str) -> list[str]:
    flags = []
    body_l = (body or "").lower()
    if contact_name and contact_name.lower() not in body_l:
        flags.append("Contact name not explicitly mentioned in body.")
    if company and company.lower() not in body_l:
        flags.append("Company name not explicitly mentioned in body.")
    generic_hits = sum(1 for pat in GENERIC_DRAFT_PATTERNS if re.search(pat, body_l))
    if generic_hits >= 2:
        flags.append("Draft appears generic; review personalization before sending.")
    return flags


def text_to_docx_bytes(title: str, body_text: str) -> bytes:
    """Wrap plain text into a simple, clean .docx (not a reproduction of the
    original resume's visual design -- just a readable, ATS-friendly draft)."""
    d = docx.Document()
    d.add_heading(title, level=1)
    for line in body_text.split("\n"):
        line = line.strip()
        if not line:
            d.add_paragraph("")
        elif line.endswith(":") and len(line) < 60:
            d.add_heading(line.rstrip(":"), level=2)
        elif line.startswith(("- ", "• ", "* ")):
            d.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            d.add_paragraph(line)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# ATS scoring engine -- part rule-based (deterministic, no AI needed), part
# AI-judged (content quality, keyword extraction). Categories sum to /100.
# ---------------------------------------------------------------------------
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s()]{7,}\d)")
URL_RE = re.compile(r"(https?://|www\.)\S+|\b\S+\.(?:com|io|dev|me|org|net)/\S*", re.I)
SECTION_PATTERNS = {
    "experience": re.compile(r"\b(experience|employment history|work experience)\b", re.I),
    "education": re.compile(r"\beducation\b", re.I),
    "skills": re.compile(r"\bskills\b", re.I),
    "summary": re.compile(r"\b(summary|objective|profile)\b", re.I),
}


def score_contact_info(text: str):
    found, score = [], 0
    if EMAIL_RE.search(text):
        score += 5; found.append("email")
    if PHONE_RE.search(text):
        score += 5; found.append("phone")
    if URL_RE.search(text):
        score += 5; found.append("link")
    return min(score, 15), found


def score_structure(text: str):
    found = [name for name, pat in SECTION_PATTERNS.items() if pat.search(text)]
    missing = [name for name in SECTION_PATTERNS if name not in found]
    score = round(len(found) / len(SECTION_PATTERNS) * 25)
    return min(score, 25), found, missing


def score_formatting(text: str):
    word_count = len(text.split())
    lines = [l for l in text.split("\n") if l.strip()]
    bullet_lines = [l for l in lines if l.strip().startswith(("-", "•", "*"))]
    bullet_ratio = (len(bullet_lines) / len(lines)) if lines else 0
    wc_score = 12 if 400 <= word_count <= 700 else (8 if 300 <= word_count <= 900 else 4)
    bullet_score = 8 if bullet_ratio >= 0.15 else (4 if bullet_ratio > 0 else 0)
    return min(wc_score + bullet_score, 20), word_count, bullet_ratio


def parse_json_loose(raw: str):
    import json
    text = re.sub(r"^```(json)?", "", raw.strip(), flags=re.I).strip()
    text = re.sub(r"```$", "", text).strip()
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r"[\{\[].*[\}\]]", text, re.S)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
        return {}


def analyze_resume_ai(resume_text: str, target_role: str, target_company: str, jd_text: str = ""):
    context = (f"Job description:\n---\n{jd_text}\n---" if jd_text.strip()
               else f"Target role: {target_role or 'a competitive role'}"
                    f"{f' at companies like {target_company}' if target_company else ''}")
    prompt = f"""You are an expert ATS resume reviewer. Analyze the resume below and respond with ONLY valid JSON \
(no markdown fences, no commentary) matching exactly this schema:

{{
  "content_quality_score": <integer 0-30>,
  "technical_keywords_found": [<specific technical/domain keywords or skills actually present in the resume, up to 20>],
  "issues": [<up to 5 short, specific problems, each under 12 words>],
  "suggestions": [<up to 6 short, specific actionable suggestions, each under 20 words>]
}}

Resume:
---
{resume_text}
---

{context}

Score content_quality_score out of 30 based on: strong action verbs, quantified achievements/impact, relevance \
and specificity (not generic filler). Do not invent facts about the candidate not present in the resume.
"""
    return parse_json_loose(call_ai(prompt, max_tokens=1200))


def run_ats_analysis(resume_text: str, target_role: str, target_company: str, jd_text: str = ""):
    contact_score, contact_found = score_contact_info(resume_text)
    structure_score, sections_found, sections_missing = score_structure(resume_text)
    formatting_score, word_count, bullet_ratio = score_formatting(resume_text)
    ai = analyze_resume_ai(resume_text, target_role, target_company, jd_text)

    content_score = max(0, min(30, int(ai.get("content_quality_score", 15) or 15)))
    keywords_found = [k for k in (ai.get("technical_keywords_found") or []) if isinstance(k, str)]
    keywords_score = min(10, round(len(keywords_found) / 15 * 10))
    overall = contact_score + structure_score + content_score + formatting_score + keywords_score

    issues = [str(i) for i in (ai.get("issues") or [])]
    for name in sections_missing:
        issues.append(f"'{name.capitalize()}' section not clearly labeled")
    if "email" not in contact_found:
        issues.append("No email address found")
    if "phone" not in contact_found:
        issues.append("No phone number found")

    suggestions = [str(s) for s in (ai.get("suggestions") or [])]
    if word_count < 400:
        suggestions.append(f"CV is {word_count} words -- aim for 400-700 words for a stronger ATS pass rate")
    elif word_count > 900:
        suggestions.append(f"CV is {word_count} words -- consider trimming toward 400-700")
    suggestions.append(f"{len(keywords_found)} technical keywords found -- try to reach 15+ for a strong ATS pass rate")

    return {
        "overall": overall,
        "categories": {
            "Contact Info": (contact_score, 15),
            "Structure": (structure_score, 25),
            "Content Quality": (content_score, 30),
            "Formatting": (formatting_score, 20),
            "Keywords": (keywords_score, 10),
        },
        "issues": issues[:8],
        "suggestions": suggestions[:8],
    }


def analyze_jd_match(resume_text: str, jd_text: str):
    prompt = f"""Extract the 10-15 most important, specific skills/qualifications/keywords from this job \
description. Respond with ONLY a JSON array of short keyword strings (1-3 words each), no commentary, e.g. \
["python","stakeholder management","aws"].

Job description:
---
{jd_text}
---
"""
    keywords = parse_json_loose(call_ai(prompt, max_tokens=400))
    if not isinstance(keywords, list):
        keywords = []
    resume_lower = resume_text.lower()
    matched = [k for k in keywords if isinstance(k, str) and k.lower().strip() in resume_lower]
    missing = [k for k in keywords if isinstance(k, str) and k.lower().strip() not in resume_lower]
    total = len(matched) + len(missing)
    pct = round(len(matched) / total * 100) if total else 0

    notes = []
    if missing:
        notes.append(f"Add these missing keywords to your CV: {', '.join(missing[:6])}")
    if total and len(matched) / total < 0.5:
        notes.append("Your CV matches less than half the JD keywords -- consider tailoring the skills section directly")
    yoe_match = re.search(r"(\d+)\+?\s*years?", jd_text, re.I)
    if yoe_match:
        notes.append(f"JD requires {yoe_match.group(1)}+ years experience -- make sure your experience duration is clearly stated")

    return {"pct": pct, "matched": matched, "missing": missing, "notes": notes}


SENIOR_TITLE_RE = re.compile(r"\b(senior|sr\.?|lead|principal|staff|director|head of|manager|vp|vice president)\b", re.I)


def suggest_role_skills(role_text: str) -> dict:
    prompt = f"""For the job role "{role_text}", suggest:
1. Up to 6 closely related alternate job titles someone might also search for.
2. Up to 10 key skills/keywords commonly required for this role.

Respond with ONLY valid JSON: {{"related_roles": [...], "skills": [...]}}. No commentary.
"""
    data = parse_json_loose(call_ai(prompt, max_tokens=400))
    if not isinstance(data, dict):
        data = {}
    return {
        "related_roles": [r for r in (data.get("related_roles") or []) if isinstance(r, str)][:8],
        "skills": [s for s in (data.get("skills") or []) if isinstance(s, str)][:12],
    }


def classify_freshers(results: list) -> set:
    """One batched AI call classifying which listings suit freshers/entry-level
    candidates, since Adzuna has no such filter natively. A title-based safety
    net removes anything obviously senior even if the AI flagged it."""
    items = [{
        "id": str(j.get("id")),
        "title": (j.get("title") or "")[:120],
        "description": (j.get("description") or "")[:300],
    } for j in results]
    if not items or not groq_key:
        return set()
    prompt = f"""Given these job listings, return ONLY a JSON array of the "id" values for listings suitable for \
freshers / entry-level candidates (0-2 years experience, new graduates, trainees). Judge from the title and \
description -- look for phrases like "entry level", "fresher", "graduate", "trainee", "0-1 years", "no experience \
required", or a clear absence of seniority requirements. Exclude anything requiring 3+ years, or titled \
Senior/Lead/Principal/Manager/Director.

Listings:
{json.dumps(items)}

Output ONLY a JSON array of matching id strings, e.g. ["123","456"]. No commentary.
"""
    try:
        ids = parse_json_loose(call_ai(prompt, max_tokens=600))
    except Exception:
        ids = []
    ai_ids = {str(i) for i in ids} if isinstance(ids, list) else set()
    return {str(j.get("id")) for j in results
            if str(j.get("id")) in ai_ids and not SENIOR_TITLE_RE.search(j.get("title") or "")}


def render_gauge(score: int, label: str = "ATS SCORE") -> str:
    score = max(0, min(100, int(score)))
    if score >= 80:
        color, tag, badge_cls = "#10b981", "Strong", "badge-ok"
    elif score >= 60:
        color, tag, badge_cls = "#e2a03f", "Average", "badge-warn"
    elif score >= 40:
        color, tag, badge_cls = "#f97316", "Needs Work", "badge-warn"
    else:
        color, tag, badge_cls = "#ef4444", "Weak", "badge-warn"
    deg = score * 3.6
    return f"""
    <div class="gauge-wrap">
      <div style="width:170px;height:170px;border-radius:50%;
           background:conic-gradient({color} {deg}deg, #ece9f9 0deg);
           display:flex;align-items:center;justify-content:center;">
        <div style="width:132px;height:132px;background:white;border-radius:50%;
             display:flex;flex-direction:column;align-items:center;justify-content:center;">
          <span style="font-size:2.4rem;font-weight:800;color:#1e1b4b;line-height:1;">{score}</span>
          <span style="color:#9ca3af;font-size:0.85rem;">/100</span>
        </div>
      </div>
      <div style="margin-top:10px;font-weight:700;color:#6b7280;letter-spacing:0.05em;font-size:0.78rem;">{label}</div>
      <span class="badge {badge_cls}" style="margin-top:6px;">{tag}</span>
    </div>
    """


def render_bar(label: str, score: int, max_score: int, color: str) -> str:
    pct = (score / max_score * 100) if max_score else 0
    return f"""
    <div class="bar-row">
      <div class="bar-label"><span>{label}</span><span style="color:{color};">{score}/{max_score}</span></div>
      <div class="bar-track"><div class="bar-fill" style="width:{pct}%;background:{color};"></div></div>
    </div>
    """


def render_chip(text: str, kind: str = "green") -> str:
    return f'<span class="chip {"chip-green" if kind == "green" else "chip-red"}">{text}</span>'


BAR_COLORS = ["#7c6ff0", "#10b981", "#e2a03f", "#3b82f6", "#ef4444"]


# ---------------------------------------------------------------------------
# 1. Resume
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# 1. Resume + ATS score
# ---------------------------------------------------------------------------
st.markdown('<div class="section-title">1. Your CV, objectively scored</div>', unsafe_allow_html=True)
st.caption("Upload your CV to get an ATS score, issues found, and suggestions -- powered by a mix of rule-based "
           "checks and AI review.")

resume_file = st.file_uploader("Upload your resume", type=["pdf", "docx", "txt"])
resume_text = None
if resume_file:
    try:
        resume_text = extract_resume_text(resume_file)
        st.success(f"Loaded {resume_file.name} ({len(resume_text)} characters extracted)")
    except Exception as e:
        st.error(f"Couldn't read that file: {e}")

col1, col2 = st.columns(2)
target_role = col1.text_input("Target role", value=st.secrets.get("TARGET_ROLE", ""))
target_company = col2.text_input("Default target company (optional -- can override per contact below)",
                                  value=st.secrets.get("TARGET_COMPANY", ""))

if st.button("Analyze CV", type="primary", disabled=not (resume_text and groq_key)):
    with st.spinner("Scoring your resume..."):
        st.session_state.ats = run_ats_analysis(resume_text, target_role, target_company)

if st.session_state.get("ats"):
    ats = st.session_state.ats
    gcol, scol = st.columns([1, 2])
    with gcol:
        st.markdown(f'<div class="score-card">{render_gauge(ats["overall"])}</div>', unsafe_allow_html=True)
    with scol:
        bars_html = "".join(
            render_bar(name, s, m, BAR_COLORS[i % len(BAR_COLORS)])
            for i, (name, (s, m)) in enumerate(ats["categories"].items())
        )
        st.markdown(f'<div class="score-card">{bars_html}</div>', unsafe_allow_html=True)

    icol, scol2 = st.columns(2)
    with icol:
        issues_html = "".join(f"<div style='margin-bottom:8px;'>❌ {i}</div>" for i in ats["issues"]) or "<div>No major issues found.</div>"
        st.markdown(f'<div class="score-card"><b style="color:#b91c1c;">⚠ ISSUES FOUND</b><div style="margin-top:12px;">{issues_html}</div></div>', unsafe_allow_html=True)
    with scol2:
        sugg_html = "".join(f"<div style='margin-bottom:8px;'>✅ {s}</div>" for s in ats["suggestions"]) or "<div>Looking solid.</div>"
        st.markdown(f'<div class="score-card"><b style="color:#5b21b6;">⚡ SUGGESTIONS</b><div style="margin-top:12px;">{sugg_html}</div></div>', unsafe_allow_html=True)

st.markdown('<div class="section-title">2. Match to a Job Description</div>', unsafe_allow_html=True)
jd_text = st.text_area(
    "Paste any job posting -- see exactly how well your CV fits",
    height=160,
    help="Also powers resume tailoring and the 'why you're a fit' pitch below.",
)

if st.button("Analyze Match", type="primary", disabled=not (resume_text and jd_text.strip() and groq_key)):
    with st.spinner("Comparing your CV to this JD..."):
        st.session_state.jd_match = analyze_jd_match(resume_text, jd_text)

if st.session_state.get("jd_match"):
    jm = st.session_state.jd_match
    pct = jm["pct"]
    badge_cls = "badge-ok" if pct >= 60 else "badge-warn"
    st.markdown(
        f'<div class="score-card"><span style="font-weight:800;font-size:1.1rem;">Match score</span>'
        f'<span class="badge {badge_cls}" style="float:right;font-size:0.95rem;padding:6px 16px;">{pct}% match</span>'
        f'<div style="clear:both;"></div></div>',
        unsafe_allow_html=True,
    )
    hcol, mcol = st.columns(2)
    with hcol:
        chips = "".join(render_chip(k, "green") for k in jm["matched"]) or "<i>None matched</i>"
        st.markdown(f'<div class="score-card"><b>YOUR CV HAS</b><div style="margin-top:10px;">{chips}</div></div>', unsafe_allow_html=True)
    with mcol:
        chips = "".join(render_chip(k, "red") for k in jm["missing"]) or "<i>Nothing missing</i>"
        st.markdown(f'<div class="score-card"><b>MISSING FROM CV</b><div style="margin-top:10px;">{chips}</div></div>', unsafe_allow_html=True)
    if jm["notes"]:
        notes_html = "".join(f"<div style='margin-bottom:8px;'>✅ {n}</div>" for n in jm["notes"])
        st.markdown(f'<div class="score-card" style="margin-top:14px;">{notes_html}</div>', unsafe_allow_html=True)

st.markdown('<div class="section-title">3. AI Tools</div>', unsafe_allow_html=True)

tab_tailor, tab_pitch, tab_cover, tab_followup, tab_translate, tab_interview = st.tabs(
    ["Tailor Resume", "Fit Pitch", "Cover Letter", "Follow-up Email", "Translate Resume", "Interview Prep"]
)

with tab_tailor:
    if st.button("Tailor resume to this JD", disabled=not (resume_text and jd_text.strip() and groq_key)):
        with st.spinner("Tailoring resume..."):
            prompt = f"""You are rewriting a candidate's resume to better match a specific job description.

Original resume:
---
{resume_text}
---

Job description:
---
{jd_text}
---

Rules -- follow these strictly:
- Only reword, reorder, re-emphasize, or tighten content that is already in the original resume.
- NEVER invent new skills, tools, employers, titles, dates, metrics, or achievements not already present or \
clearly implied by the original.
- Prioritize and surface the experience most relevant to this JD nearer the top of each section.
- Naturally work in JD keywords/phrasing only where they genuinely match real content.
- Keep the same overall sections (e.g. Experience, Education, Skills) the original has.
- Output the full tailored resume as plain text, ready to review -- no commentary, no markdown formatting, \
just the resume content with clear section headers.
"""
            st.session_state.tailored_resume = call_ai(prompt, max_tokens=2500)

    if st.session_state.get("tailored_resume"):
        st.warning("Content-only draft (plain formatting) -- doesn't preserve your original resume's visual "
                   "design. Copy what's useful back into your real resume, or download this version.")
        st.text_area("Tailored resume text", value=st.session_state.tailored_resume, height=300, key="tailored_view")
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes("Tailored Resume", st.session_state.tailored_resume),
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_tailored",
        )

with tab_pitch:
    if st.button("Generate 'why you're a fit' pitch", disabled=not (resume_text and groq_key)):
        with st.spinner("Writing pitch..."):
            prompt = f"""Write exactly 3 short sentences (not bullet points) explaining why this candidate is a \
strong fit for this specific role. This will be used as an opening pitch in outreach emails or pasted into an \
application's free-text "why are you interested" field.

Candidate resume:
---
{resume_text}
---

{"Job description:\n---\n" + jd_text + "\n---" if jd_text.strip() else f"Target role: {target_role or 'a relevant role'}"}

Only reference real, specific experience from the resume -- no generic claims like "I am a hard worker." \
Output only the 3 sentences, nothing else.
"""
            st.session_state.fit_pitch = call_ai(prompt, max_tokens=250)

    if st.session_state.get("fit_pitch"):
        st.text_area("Copy this into emails or application forms", value=st.session_state.fit_pitch,
                      height=100, key="fit_pitch_view")

with tab_cover:
    st.caption("A full, ready-to-send cover letter -- uses the job description above if you've pasted one.")
    if st.button("Generate cover letter", disabled=not (resume_text and groq_key)):
        with st.spinner("Writing cover letter..."):
            prompt = f"""Write a complete, professional cover letter for this candidate.

Candidate resume:
---
{resume_text}
---

{"Job description:\n---\n" + jd_text + "\n---" if jd_text.strip() else f"Target role: {target_role or 'a relevant role'}{f' at {target_company}' if target_company else ''}"}
Sender's name to sign as: {sender_name or "[Your Name]"}

Requirements:
- 250-350 words, 3-4 paragraphs: an opening hook, 1-2 body paragraphs connecting specific real resume experience \
to the role's needs, a closing call to action.
- Professional but not stiff -- avoid cliches like "I am writing to express my interest."
- Only reference real experience from the resume -- never invent achievements, employers, or skills.
- Output the letter only, no commentary, ready to copy-paste or export.
"""
            st.session_state.cover_letter = call_ai(prompt, max_tokens=900)

    if st.session_state.get("cover_letter"):
        st.text_area("Cover letter", value=st.session_state.cover_letter, height=320, key="cover_letter_view")
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes("Cover Letter", st.session_state.cover_letter),
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_cover",
        )

with tab_followup:
    st.caption("For when you've already applied or emailed and haven't heard back.")
    fcol1, fcol2 = st.columns(2)
    fu_company = fcol1.text_input("Company", value=target_company, key="fu_company")
    fu_role = fcol2.text_input("Role applied for", value=target_role, key="fu_role")
    fcol3, fcol4 = st.columns(2)
    fu_contact_name = fcol3.text_input("Contact name (optional)", key="fu_contact_name")
    fu_days = fcol4.number_input("Days since you applied/last contacted", min_value=1, value=7, key="fu_days")
    fu_notes = st.text_input("Anything else to mention (optional)", key="fu_notes")

    if st.button("Generate follow-up email", disabled=not (resume_text and groq_key and fu_company and fu_role)):
        with st.spinner("Drafting follow-up..."):
            prompt = f"""Write a short, polite follow-up email checking on a job application.

Candidate resume (for context):
---
{resume_text}
---

Company: {fu_company}
Role applied for: {fu_role}
Contact: {fu_contact_name or "Hiring team"}
It has been {fu_days} days since applying/last contact.
Extra context to weave in if useful: {fu_notes or "none"}
Sender's name to sign as: {sender_name or "[Your Name]"}

Requirements:
- Subject line: short, references the role.
- Body: 60-100 words. Polite, brief, no guilt-tripping, reaffirms interest with one concrete detail from the \
resume, ends with a low-pressure ask (e.g. any update, happy to provide more info).
- Output format exactly:
SUBJECT: <subject line>
BODY:
<email body>
"""
            text = call_ai(prompt, max_tokens=350)
            subject, body = "", text
            if "SUBJECT:" in text and "BODY:" in text:
                subject = text.split("SUBJECT:", 1)[1].split("BODY:", 1)[0].strip()
                body = text.split("BODY:", 1)[1].strip()
            st.session_state.followup = {"subject": subject, "body": body}

    if st.session_state.get("followup"):
        fu = st.session_state.followup
        fu["subject"] = st.text_input("Subject", value=fu["subject"], key="fu_subject_view")
        fu["body"] = st.text_area("Body", value=fu["body"], height=180, key="fu_body_view")
        fu_send_to = st.text_input("Send to (email address)", key="fu_send_to")
        if st.button("Send this follow-up now", disabled=not (fu_send_to and resume_file and gmail_address and gmail_app_password)):
            try:
                send_email(fu_send_to, fu["subject"], fu["body"], resume_file.getvalue(), resume_file.name)
                st.success(f"Sent to {fu_send_to}")
            except Exception as e:
                st.error(f"Failed to send: {e}")

with tab_translate:
    st.caption("Content-only translation -- review with a native speaker before using for something high-stakes.")
    lang = st.selectbox("Translate resume to", [
        "Spanish", "French", "German", "Portuguese", "Italian", "Dutch",
        "Hindi", "Mandarin Chinese", "Japanese", "Korean", "Arabic",
    ], key="translate_lang")
    if st.button("Translate resume", disabled=not (resume_text and groq_key)):
        with st.spinner(f"Translating to {lang}..."):
            prompt = f"""Translate this resume into {lang}. Keep the same structure and section order. Use \
natural, professional resume phrasing appropriate for {lang}-speaking employers, not a literal word-for-word \
translation. Do not add, remove, or change any facts, dates, numbers, or claims. Output only the translated \
resume text, no commentary.

Resume:
---
{resume_text}
---
"""
            st.session_state.translated_resume = call_ai(prompt, max_tokens=2500)

    if st.session_state.get("translated_resume"):
        st.text_area(f"Translated resume ({lang})", value=st.session_state.translated_resume, height=320, key="translated_view")
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes(f"Resume ({lang})", st.session_state.translated_resume),
            file_name=f"resume_{lang.lower().replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_translated",
        )

with tab_interview:
    st.caption("Practice prep -- questions and talking points to prepare *before* an interview, based on your "
               "real experience. (Not a live tool for feeding you answers during an actual interview -- that "
               "would mean the interviewer is evaluating answers you didn't actually give.)")
    if st.button("Generate interview prep sheet", disabled=not (resume_text and groq_key)):
        with st.spinner("Preparing questions..."):
            prompt = f"""Create an interview prep sheet for this candidate.

Candidate resume:
---
{resume_text}
---

{"Job description:\n---\n" + jd_text + "\n---" if jd_text.strip() else f"Target role: {target_role or 'a relevant role'}"}

Produce:
1. 5 likely behavioral questions for this role, each with a brief note on which real resume experience to draw \
on and a suggested STAR-structure outline (Situation/Task/Action/Result) -- do not write out full answers, just \
the outline and which real experience fits, so the candidate practices in their own words.
2. 5 likely technical/role-specific questions based on the JD or target role.
3. 3 smart questions the candidate can ask the interviewer.

Only reference real experience already in the resume. No invented achievements.
"""
            st.session_state.interview_prep = call_ai(prompt, max_tokens=1500)

    if st.session_state.get("interview_prep"):
        st.markdown(st.session_state.interview_prep)
        st.download_button(
            "Download as .docx",
            data=text_to_docx_bytes("Interview Prep Sheet", st.session_state.interview_prep),
            file_name="interview_prep.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_interview",
        )


# ---------------------------------------------------------------------------
# 4. Find open roles (legitimate public job API -- not scraping LinkedIn/Indeed/Naukri)
# ---------------------------------------------------------------------------
st.markdown('<div class="section-title">4. Find open roles</div>', unsafe_allow_html=True)
st.caption("Searches via Adzuna's public API (aggregates many boards legitimately, official API -- no scraping).")

if not (adzuna_app_id and adzuna_app_key):
    st.info("Add `ADZUNA_APP_ID` and `ADZUNA_APP_KEY` to secrets to enable this -- "
            "https://developer.adzuna.com/")
else:
    st.session_state.setdefault("job_search_kw", target_role)
    st.session_state.setdefault("my_skills", [])

    scol1, scol2, scol3 = st.columns([2, 1, 1])
    search_keyword = scol1.text_input("Role", key="job_search_kw")
    search_country = scol2.selectbox("Country", ["us", "gb", "in", "ca", "au", "de", "fr"], index=0)
    search_page = scol3.number_input("Page", min_value=1, value=1)

    if st.button("Suggest related roles & skills", disabled=not (search_keyword and groq_key)):
        with st.spinner("Thinking..."):
            st.session_state.role_suggestions = suggest_role_skills(search_keyword)

    rs = st.session_state.get("role_suggestions")
    if rs and (rs.get("related_roles") or rs.get("skills")):
        if rs.get("related_roles"):
            st.caption("Related roles -- click to search that instead:")
            rcols = st.columns(min(len(rs["related_roles"]), 6) or 1)
            for i, r in enumerate(rs["related_roles"]):
                if rcols[i % len(rcols)].button(r, key=f"relrole_{i}", width='stretch'):
                    st.session_state.job_search_kw = r
                    st.rerun()
        if rs.get("skills"):
            st.caption("Suggested skills -- click to add to your skill list:")
            scols = st.columns(min(len(rs["skills"]), 6) or 1)
            for i, s in enumerate(rs["skills"]):
                if scols[i % len(scols)].button(f"+ {s}", key=f"addskill_{i}", width='stretch'):
                    if s not in st.session_state.my_skills:
                        st.session_state.my_skills.append(s)
                    st.rerun()

    if st.session_state.my_skills:
        chips = "".join(f'<span class="skill-chip">{s}</span>' for s in st.session_state.my_skills)
        st.markdown(f"**Your skills:** {chips}", unsafe_allow_html=True)
        if st.button("Clear skills"):
            st.session_state.my_skills = []
            st.rerun()

    fcol1, fcol2 = st.columns(2)
    fresher_only = fcol1.checkbox("🎓 Freshers / entry-level only",
                                   help="AI reads each listing's title & description to judge this -- Adzuna has no native filter for it.")
    skill_filter_on = fcol2.checkbox("Only listings matching my skills", disabled=not st.session_state.my_skills)

    if st.button("Search open roles", type="primary"):
        with st.spinner("Searching..."):
            try:
                resp = requests.get(
                    f"https://api.adzuna.com/v1/api/jobs/{search_country}/search/{int(search_page)}",
                    params={
                        "app_id": adzuna_app_id,
                        "app_key": adzuna_app_key,
                        "what": search_keyword,
                        "content-type": "application/json",
                    },
                    timeout=15,
                )
                resp.raise_for_status()
                st.session_state.job_results = resp.json().get("results", [])
            except Exception as e:
                st.error(f"Search failed: {e}")
                st.session_state.job_results = []
        if st.session_state.job_results and groq_key:
            with st.spinner("Checking which listings suit freshers..."):
                st.session_state.fresher_ids = classify_freshers(st.session_state.job_results)
        else:
            st.session_state.fresher_ids = set()

    MANUAL_ONLY_DOMAINS = ["linkedin.com", "indeed.com", "naukri.com", "glassdoor.com"]

    results = st.session_state.get("job_results", [])
    if results:
        st.session_state.setdefault("saved_jobs", set())
        fresher_ids = st.session_state.get("fresher_ids", set())
        my_skills_lower = [s.lower() for s in st.session_state.my_skills]

        def skill_matches(job):
            text = f"{job.get('title', '')} {job.get('description', '')}".lower()
            return [s for s in st.session_state.my_skills if s.lower() in text]

        display_results = []
        for j in results:
            if fresher_only and str(j.get("id")) not in fresher_ids:
                continue
            matches = skill_matches(j) if my_skills_lower else []
            if skill_filter_on and my_skills_lower and not matches:
                continue
            display_results.append((j, matches))

        fcol1, fcol2 = st.columns([3, 1])
        fcol1.markdown(f"**Job Listings** &nbsp; <span style='color:var(--text-muted);'>{len(display_results)} "
                       f"of {len(results)} results</span>", unsafe_allow_html=True)
        show_saved_only = fcol2.checkbox("Saved only")
        if show_saved_only:
            display_results = [(j, m) for j, m in display_results if j.get("id") in st.session_state.saved_jobs]

        cols = st.columns(3)
        for idx, (job, matches) in enumerate(display_results):
            url = job.get("redirect_url", "")
            company = (job.get("company") or {}).get("display_name", "Unknown company")
            title = job.get("title", "Untitled role")
            location = (job.get("location") or {}).get("display_name", "")
            job_id = job.get("id")
            manual_only = any(d in url for d in MANUAL_ONLY_DOMAINS)
            badges = (f'<span class="badge badge-warn">MANUAL APPLY</span>' if manual_only
                      else f'<span class="badge badge-ok">AUTO-APPLY OK</span>')
            if str(job_id) in fresher_ids:
                badges += ' <span class="badge badge-ok">🎓 FRESHER FRIENDLY</span>'
            if st.session_state.my_skills:
                badges += f' <span class="badge">{len(matches)}/{len(st.session_state.my_skills)} skills matched</span>'

            with cols[idx % 3]:
                st.markdown(
                    f"""<div class="job-card">
                        <div style="font-weight:800;font-size:1.05rem;">{title}</div>
                        <div style="opacity:0.85;margin:4px 0;">🏢 {company}</div>
                        <div style="opacity:0.7;font-size:0.9rem;margin-bottom:8px;">📍 {location}</div>
                        {badges}
                    </div>""",
                    unsafe_allow_html=True,
                )
                bcol1, bcol2 = st.columns(2)
                bcol1.link_button("Apply Now →", url, width='stretch')
                is_saved = job_id in st.session_state.saved_jobs
                if bcol2.button("★ Saved" if is_saved else "☆ Save", key=f"save_{job_id}_{idx}", width='stretch'):
                    if is_saved:
                        st.session_state.saved_jobs.discard(job_id)
                    else:
                        st.session_state.saved_jobs.add(job_id)
                    st.rerun()
                if not manual_only:
                    if st.button("+ Add to auto-apply list", key=f"add_{job_id}_{idx}", width='stretch'):
                        st.session_state.setdefault("career_urls", [])
                        if url not in st.session_state.career_urls:
                            st.session_state.career_urls.append(url)
                        st.success("Added below ↓")
                else:
                    st.caption("This platform prohibits auto-apply bots -- use the fit-pitch/tailored resume above, then apply yourself.")

    if st.session_state.get("career_urls"):
        st.subheader("Your auto-apply list (for apply_bot.py)")
        urls_text = "\n".join(st.session_state.career_urls)
        st.text_area("One URL per line -- copy into career_urls.txt", value=urls_text, height=100)
        st.download_button("Download career_urls.txt", data=urls_text, file_name="career_urls.txt")

# ---------------------------------------------------------------------------
# 5. Contacts
# ---------------------------------------------------------------------------
st.markdown('<div class="section-title">5. Contacts</div>', unsafe_allow_html=True)
st.caption("Add each person you want to email. Leave name/role/job_title blank to fall back to defaults. Optional "
           "personalization columns: company_note, contact_note, project_match, tone, custom_subject.")

upload_col, _ = st.columns([1, 2])
uploaded_contacts = upload_col.file_uploader(
    "Or upload a contact list (.csv or .xlsx)",
    type=["csv", "xlsx"],
    help="Needs 'email'. Optional: name, role, company, job_title, company_note, contact_note, project_match, tone, custom_subject.",
)
if uploaded_contacts is not None:
    try:
        if uploaded_contacts.name.lower().endswith(".xlsx"):
            up_df = pd.read_excel(uploaded_contacts)
        else:
            up_df = pd.read_csv(uploaded_contacts)
        up_df.columns = [c.strip().lower() for c in up_df.columns]
        missing_required = [c for c in CONTACT_REQUIRED_COLUMNS if c not in up_df.columns]
        if missing_required:
            st.error(f"Missing required column(s): {', '.join(missing_required)}")
        else:
            for col in CONTACT_OPTIONAL_COLUMNS:
                if col not in up_df.columns:
                    up_df[col] = ""
            up_df = up_df[["email"] + CONTACT_OPTIONAL_COLUMNS].fillna("")
            up_df = up_df.astype(str).apply(lambda col: col.str.strip())

            blank_rows, invalid_rows, duplicate_emails = validate_uploaded_contacts(up_df)
            if blank_rows:
                st.warning(f"Rows with blank email: {blank_rows[:10]}{' ...' if len(blank_rows) > 10 else ''}")
            if invalid_rows:
                st.warning(f"Rows with invalid email format: {invalid_rows[:10]}{' ...' if len(invalid_rows) > 10 else ''}")
            if duplicate_emails:
                st.warning("Duplicate email(s) detected: " + ", ".join(duplicate_emails[:10]) +
                           (" ..." if len(duplicate_emails) > 10 else ""))

            st.session_state.recipients_df = up_df
            st.success(f"Loaded {len(up_df)} contact(s) from {uploaded_contacts.name}")
            st.caption("Parsed preview (first 3 rows)")
            st.dataframe(up_df.head(3), use_container_width=True, hide_index=True)
    except Exception as e:
        st.error(f"Couldn't read that file: {e}")

if "recipients_df" not in st.session_state:
    st.session_state.recipients_df = pd.DataFrame(
        [{
            "email": "", "name": "", "role": "", "company": "", "job_title": "",
            "company_note": "", "contact_note": "", "project_match": "", "tone": "", "custom_subject": ""
        }]
    )

recipients_df = st.data_editor(
    st.session_state.recipients_df,
    num_rows="dynamic",
    width='stretch',
    key="recipients_editor",
)
st.session_state.recipients_df = recipients_df

# ---------------------------------------------------------------------------
# 5. Draft
# ---------------------------------------------------------------------------
st.markdown('<div class="section-title">6. Draft emails</div>', unsafe_allow_html=True)

global_fixed_message = st.text_area(
    "About me / what I’m looking for (included in every email)",
    value=(st.session_state.get("global_fixed_message")
           or f"I’m looking for {target_role or 'a relevant role'} opportunities and would value a quick conversation."),
    height=90,
)
st.session_state.global_fixed_message = global_fixed_message

cta_options = ["referral request", "recruiter connect", "hiring manager intro"]
current_cta = st.session_state.get("cta_style", "recruiter connect")
if current_cta not in cta_options:
    current_cta = "recruiter connect"
cta_style = st.selectbox(
    "Call to action style",
    options=cta_options,
    index=cta_options.index(current_cta),
)
st.session_state.cta_style = cta_style

lcol1, lcol2 = st.columns(2)
lock_fixed_section = lcol1.checkbox("Lock fixed message text (verbatim)", value=True)
lock_personalization_section = lcol2.checkbox("Lock row personalization snippets (verbatim)", value=True)

st.caption("Template preview")
st.info(
    "Greeting → Fixed message → Company/contact personalization → Fit proof from your resume → "
    "CTA → Signoff"
)

if st.button("Generate drafts", disabled=not (resume_text and groq_key)):
    rows = recipients_df[recipients_df["email"].astype(str).str.strip() != ""]
    if rows.empty:
        st.warning("Add at least one contact with an email address above.")
    else:
        drafts = []
        cta_text_map = {
            "referral request": "Ask for a referral if they think your background fits.",
            "recruiter connect": "Ask for a brief recruiter conversation about relevant roles.",
            "hiring manager intro": "Ask for a short intro call with the hiring manager/team.",
        }
        fixed_message = global_fixed_message.strip() or (
            f"I’m looking for {target_role or 'a relevant role'} opportunities and would value a quick conversation."
        )
        progress = st.progress(0.0, text="Drafting emails...")
        for i, row in enumerate(rows.to_dict("records")):
            company = row.get("company") or target_company
            job_title = row.get("job_title") or target_role
            contact_name = row.get("name") or name_from_email(row["email"]) or "Hiring contact"
            company_note = (row.get("company_note") or "").strip()
            contact_note = (row.get("contact_note") or "").strip()
            project_match = (row.get("project_match") or "").strip()
            tone = (row.get("tone") or "").strip() or "professional"
            custom_subject = (row.get("custom_subject") or "").strip()

            personalization_lines = []
            if company_note:
                personalization_lines.append(f"- Company note: {company_note}")
            if contact_note:
                personalization_lines.append(f"- Contact note: {contact_note}")
            if project_match:
                personalization_lines.append(f"- Project/skill match to mention: {project_match}")
            personalization_block = "\n".join(personalization_lines) if personalization_lines else "None provided."
            lock_fixed_instruction = ("Include the fixed message EXACTLY as written, without rephrasing."
                                      if lock_fixed_section else "You may lightly polish the fixed message.")
            lock_personalization_instruction = (
                "If personalization snippets are present, keep their meaning and key wording intact."
                if lock_personalization_section else
                "You may rewrite personalization snippets for flow."
            )
            prompt = f"""Write a concise, professional job-outreach email from a candidate to a contact at a company.

Candidate resume (for context on real background -- do not fabricate anything not in here):
---
{resume_text}
---
{f"Job description for the target role:\n---\n{jd_text}\n---\n" if jd_text.strip() else ""}
Recipient: {contact_name}{f", {row.get('role')}" if row.get("role") else ""} at {company or "the company"}
Target position: {job_title or "a relevant open role"}
Sender's name to sign as: {sender_name or "[Your Name]"}
Preferred tone: {tone}
CTA style: {cta_style} ({cta_text_map.get(cta_style, '')})

Fixed message to include in every email:
<<FIXED_MESSAGE>>
{fixed_message}
<</FIXED_MESSAGE>>

Per-contact personalization snippets:
{personalization_block}

Requirements:
- Subject line: short, specific, not spammy
- Body: 120-180 words, warm but professional, no generic filler ("I am writing to express...")
- Reference 1-2 concrete, real details from the resume that make the candidate relevant{" to the job description" if jd_text.strip() else ""}
- Must mention both the recipient/company context and the target role
- Must include the fixed message block in the body
- Use personalization snippets when present
- Mention the resume is attached
- End with a clear, low-pressure call to action matching CTA style
- {lock_fixed_instruction}
- {lock_personalization_instruction}
- No markdown, no placeholders left unfilled except sign-off name if not provided
- Output format exactly:
SUBJECT: <subject line>
BODY:
<email body>
"""
            text = call_ai(prompt, max_tokens=600)
            subject, body = "", text
            if "SUBJECT:" in text and "BODY:" in text:
                subject = text.split("SUBJECT:", 1)[1].split("BODY:", 1)[0].strip()
                body = text.split("BODY:", 1)[1].strip()
            if custom_subject:
                subject = custom_subject
            quality_flags = evaluate_draft_quality(body, contact_name, company or "")
            drafts.append({
                "email": row["email"],
                "subject": subject,
                "body": body,
                "quality_flags": quality_flags,
            })
            progress.progress((i + 1) / len(rows), text=f"Drafted {i + 1}/{len(rows)}")
        st.session_state.drafts = drafts
        st.rerun()

if st.session_state.get("drafts"):
    st.subheader("Review & edit before sending")
    flagged = sum(1 for d in st.session_state.drafts if d.get("quality_flags"))
    if flagged:
        st.warning(f"{flagged} draft(s) flagged for manual review due to weak personalization.")
    for i, d in enumerate(st.session_state.drafts):
        with st.expander(f"✉️ {d['email']} — {d['subject']}"):
            if d.get("quality_flags"):
                st.warning(" | ".join(d["quality_flags"]))
            d["subject"] = st.text_input("Subject", value=d["subject"], key=f"subj_{i}")
            d["body"] = st.text_area("Body", value=d["body"], height=220, key=f"body_{i}")

# ---------------------------------------------------------------------------
# 7. Send
# ---------------------------------------------------------------------------
st.markdown('<div class="section-title">7. Send</div>', unsafe_allow_html=True)

ready_to_send = bool(
    st.session_state.get("drafts") and resume_file and gmail_address and gmail_app_password
)
if not ready_to_send:
    st.info("Generate drafts above and fill in your Gmail credentials in the sidebar to enable sending.")

confirm = st.text_input('Type SEND (all caps) to enable the send button', value="")
send_clicked = st.button(
    f"Send {len(st.session_state.get('drafts', []))} email(s) now",
    disabled=not (ready_to_send and confirm == "SEND"),
    type="primary",
)

if send_clicked:
    resume_bytes = resume_file.getvalue()
    results = []
    progress = st.progress(0.0, text="Sending...")
    drafts = st.session_state.drafts
    for i, d in enumerate(drafts):
        try:
            send_email(d["email"], d["subject"], d["body"], resume_bytes, resume_file.name)
            results.append({"email": d["email"], "status": "sent"})
        except Exception as e:
            results.append({"email": d["email"], "status": f"FAILED: {e}"})
        progress.progress((i + 1) / len(drafts), text=f"Sent {i + 1}/{len(drafts)}")
        if i < len(drafts) - 1:
            time.sleep(seconds_between)
    st.success("Done.")
    st.table(pd.DataFrame(results))
